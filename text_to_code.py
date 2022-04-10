import docx

doc = docx.Document("C:\withpython\Catcha_word.docx")
all_paras = doc.paragraphs
mydoc = docx.Document()
####################################################################
for para in all_paras:
	tex = para.text
	#bắt đầu 
	vt = tex.find('\\')
	vt += 1

	#tách đáp án từ vtdapan -> vt1dapan  
	vt1 = vt
	while(vt1 < len(tex)):
		if (tex[vt1] == "\\"):
			break
		vt1+=1
	vtdapan = vt
	vt1dapan = vt1

	#tách câu hỏi 
	vt1+=1

	ghi=mydoc.add_paragraph("elif(tex == \"")

	while(vt1 < len(tex)):
		if (tex[vt1] != "/"):
			ghi.add_run(tex[vt1])
		elif (tex[vt1] != "\n"):
			pass
		else:
			ghi.add_run(" ")
		vt1 += 1
	ghi.add_run(" ")
	ghi.add_run("\"):")

	mydoc.add_paragraph("\treturn ").add_run(tex[vtdapan:vt1dapan])
	#kết thúc

#################################################
mydoc.save("E:/my_written_file.docx")
